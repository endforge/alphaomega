"""
File: test_text_extractor_formats.py

Purpose:
    Verify TextExtractor behavior for each currently supported
    extraction format without accessing Microsoft Graph or any
    external Source of Truth.

Tests:
    - Plain text extraction.
    - DOCX extraction.
    - PDF extraction.
    - XLSX extraction.
    - HTML extraction.
    - Unsupported format rejection.
"""

import io

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from scripts.extraction.text_extractor import (
    TextExtractor,
)


def build_docx_bytes():
    """
    Create a synthetic DOCX document in memory.
    """

    buffer = io.BytesIO()

    document = Document()

    document.add_paragraph(
        "AlphaOmega DOCX test."
    )

    document.add_paragraph(
        "Second paragraph."
    )

    document.save(buffer)

    return buffer.getvalue()


def build_pdf_bytes():
    """
    Create a synthetic PDF document in memory.
    """

    buffer = io.BytesIO()

    pdf = canvas.Canvas(buffer)

    pdf.drawString(
        72,
        720,
        "AlphaOmega PDF test."
    )

    pdf.drawString(
        72,
        700,
        "Second PDF line."
    )

    pdf.save()

    return buffer.getvalue()


def build_xlsx_bytes():
    """
    Create a synthetic XLSX workbook in memory.
    """

    buffer = io.BytesIO()

    workbook = Workbook()

    sheet = workbook.active
    sheet.title = "AlphaOmega"

    sheet["A1"] = "Object"
    sheet["B1"] = "Status"

    sheet["A2"] = "Bogmire"
    sheet["B2"] = "Canonical"

    workbook.save(buffer)

    return buffer.getvalue()


def test_plain_text():
    """
    Verify plain UTF-8 text extraction.
    """

    result = TextExtractor.extract(
        "alphaomega.txt",
        b"AlphaOmega plain text.",
    )

    assert (
        result
        == "AlphaOmega plain text."
    )

    print(
        "PASS: Plain text extraction correct."
    )


def test_docx():
    """
    Verify DOCX paragraph extraction.
    """

    result = TextExtractor.extract(
        "alphaomega.docx",
        build_docx_bytes(),
    )

    assert (
        "AlphaOmega DOCX test."
        in result
    )

    assert (
        "Second paragraph."
        in result
    )

    print(
        "PASS: DOCX extraction correct."
    )


def test_pdf():
    """
    Verify PDF page text extraction.
    """

    result = TextExtractor.extract(
        "alphaomega.pdf",
        build_pdf_bytes(),
    )

    assert (
        "AlphaOmega PDF test."
        in result
    )

    assert (
        "Second PDF line."
        in result
    )

    print(
        "PASS: PDF extraction correct."
    )


def test_xlsx():
    """
    Verify XLSX worksheet and populated-cell extraction.
    """

    result = TextExtractor.extract(
        "alphaomega.xlsx",
        build_xlsx_bytes(),
    )

    assert (
        "--- Sheet: AlphaOmega ---"
        in result
    )

    assert (
        "Object | Status"
        in result
    )

    assert (
        "Bogmire | Canonical"
        in result
    )

    print(
        "PASS: XLSX extraction correct."
    )


def test_html():
    """
    Verify visible HTML text extraction.
    """

    html = (
        b"<html>"
        b"<body>"
        b"<h1>AlphaOmega HTML test.</h1>"
        b"<p>Canonical page content.</p>"
        b"</body>"
        b"</html>"
    )

    result = TextExtractor.extract(
        "alphaomega.html",
        html,
    )

    assert (
        "AlphaOmega HTML test."
        in result
    )

    assert (
        "Canonical page content."
        in result
    )

    print(
        "PASS: HTML extraction correct."
    )


def test_unsupported_format():
    """
    Verify unsupported file formats are explicitly rejected.
    """

    try:
        TextExtractor.extract(
            "alphaomega.jpg",
            b"synthetic image bytes",
        )

        raise AssertionError(
            "Unsupported format was accepted."
        )

    except ValueError as error:
        assert (
            "Unsupported extraction format"
            in str(error)
        )

    print(
        "PASS: Unsupported format rejected."
    )


def main():
    """
    Run format-specific TextExtractor tests.
    """

    print(
        "\nRunning TextExtractor format tests...\n"
    )

    test_plain_text()
    test_docx()
    test_pdf()
    test_xlsx()
    test_html()
    test_unsupported_format()

    print(
        "\nTextExtractor format tests PASSED.\n"
    )


if __name__ == "__main__":
    main()